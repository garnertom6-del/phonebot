"""Shared DOCX construction for the CARF packet builder.

Renders a small, predictable subset of Markdown into a styled Word document:
  # / ## / ###   headings
  **bold**       inline bold
  - item         bullet
  1. item        numbered
  | a | b |      table (a |---| row marks the header)
  ---            horizontal rule
Anything else becomes a body paragraph.
"""
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY = RGBColor(0x1F, 0x3B, 0x63)
GREY = RGBColor(0x55, 0x55, 0x55)
RED = RGBColor(0xA8, 0x1C, 0x1C)


def new_document():
    doc = Document()
    for s in doc.sections:
        s.top_margin = s.bottom_margin = Inches(0.7)
        s.left_margin = s.right_margin = Inches(0.8)
    st = doc.styles["Normal"]
    st.font.name = "Calibri"
    st.font.size = Pt(10.5)
    st.paragraph_format.space_after = Pt(5)
    return doc


def _set_cell_bg(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hexcolor)
    tcPr.append(shd)


def add_bold_runs(par, text):
    """Split on **bold** and emit runs."""
    for i, chunk in enumerate(re.split(r"\*\*(.+?)\*\*", text)):
        if not chunk:
            continue
        run = par.add_run(chunk)
        run.bold = i % 2 == 1
    return par


def heading(doc, text, level=1, color=NAVY):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    r.bold = True
    r.font.color.rgb = color
    r.font.size = Pt({1: 16, 2: 13, 3: 11.5}.get(level, 11))
    return p


def rule(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:color"), "1F3B63")
    pbdr.append(bottom)
    pPr.append(pbdr)
    return p


def notice(doc, text, color=RED):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    r = p.add_run(text)
    r.bold = True
    r.font.color.rgb = color
    r.font.size = Pt(10.5)
    return p


def page_break(doc):
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def _table(doc, rows):
    ncol = max(len(r) for r in rows)
    t = doc.add_table(rows=0, cols=ncol)
    t.style = "Table Grid"
    t.autofit = True
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for ci in range(ncol):
            txt = row[ci] if ci < len(row) else ""
            cell = cells[ci]
            cell.text = ""
            par = cell.paragraphs[0]
            par.paragraph_format.space_after = Pt(1)
            add_bold_runs(par, txt)
            for r in par.runs:
                r.font.size = Pt(8.5)
                if ri == 0:
                    r.bold = True
            if ri == 0:
                _set_cell_bg(cell, "E8EDF5")
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


def render_markdown(doc, md, base_level=1):
    """Render the supported Markdown subset into doc."""
    lines = md.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        if stripped.startswith("|"):
            block = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                block.append(lines[i].strip())
                i += 1
            rows = []
            for b in block:
                if re.fullmatch(r"\|[\s:|-]+\|", b):
                    continue
                cells = [c.strip() for c in b.strip("|").split("|")]
                rows.append(cells)
            if rows:
                _table(doc, rows)
            continue

        if stripped in ("---", "***", "___"):
            rule(doc)
            i += 1
            continue

        m = re.match(r"^(#{1,4})\s+(.*)$", stripped)
        if m:
            heading(doc, m.group(2), min(len(m.group(1)) + base_level - 1, 3))
            i += 1
            continue

        m = re.match(r"^[-*]\s+(.*)$", stripped)
        if m:
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(2)
            add_bold_runs(p, m.group(1))
            i += 1
            continue

        m = re.match(r"^(\d+)\.\s+(.*)$", stripped)
        if m:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            p.paragraph_format.space_after = Pt(2)
            add_bold_runs(p, f"{m.group(1)}. {m.group(2)}")
            i += 1
            continue

        p = doc.add_paragraph()
        add_bold_runs(p, stripped)
        i += 1


def title_page(doc, agency, doc_title, subtitle, meta_lines, warning=None):
    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(agency)
    r.bold = True
    r.font.size = Pt(24)
    r.font.color.rgb = NAVY

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(doc_title)
    r.bold = True
    r.font.size = Pt(18)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(subtitle)
    r.font.size = Pt(12)
    r.font.color.rgb = GREY

    doc.add_paragraph()
    if warning:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(warning)
        r.bold = True
        r.font.size = Pt(11)
        r.font.color.rgb = RED

    for _ in range(3):
        doc.add_paragraph()
    for line in meta_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(line)
        r.font.size = Pt(10)
        r.font.color.rgb = GREY
    page_break(doc)


def footer_text(doc, text):
    for section in doc.sections:
        p = section.footer.paragraphs[0]
        p.text = text
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.font.size = Pt(8)
            r.font.color.rgb = GREY
